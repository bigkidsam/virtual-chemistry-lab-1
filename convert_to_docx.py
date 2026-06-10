import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_inline_styles(paragraph, text):
    # Match bold italic (***), bold (**), italic (*), and inline code (`)
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.0)
            run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B) # Crimson red for code
        else:
            paragraph.add_run(part)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E5E7EB"/></w:pBdr>')
    pPr.append(pBdr)

def add_heading(doc, text, level, is_cover=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI Semibold'
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy
        if is_cover:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI Semibold'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        if is_cover:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI Semibold'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal
    elif level == 4:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI Semibold'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

def add_bullet_item(doc, text):
    cleaned = re.sub(r'^[\-\*]\s+', '', text)
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    parse_inline_styles(p, cleaned)

def add_numbered_item(doc, text):
    cleaned = re.sub(r'^\d+\.\s+', '', text)
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    parse_inline_styles(p, cleaned)

def flush_paragraph_buffer(doc, buffer):
    if not buffer:
        return
    text = " ".join([line.strip() for line in buffer])
    if text:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        parse_inline_styles(p, text)
    buffer.clear()

def process_table(doc, table_lines):
    rows_data = []
    for line in table_lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\|[\s\-\:\|]+$', line):
            continue
        parts = line.split('|')
        if parts[0] == '':
            parts = parts[1:]
        if parts and parts[-1] == '':
            parts = parts[:-1]
        
        cells = [c.strip() for c in parts]
        rows_data.append(cells)
        
    if not rows_data:
        return
        
    num_cols = max(len(row) for row in rows_data)
    for row in rows_data:
        while len(row) < num_cols:
            row.append('')
            
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    set_table_borders(table)
    
    for row_idx, row_cells in enumerate(rows_data):
        trPr = table.rows[row_idx]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        if row_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for col_idx, text in enumerate(row_cells):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            
            if row_idx == 0:
                set_cell_background(cell, "1E3A8A")  # Navy
                parse_inline_styles(p, text)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9.5)
                set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
            else:
                if row_idx % 2 == 1:
                    set_cell_background(cell, "F9FAFB")  # Zebra striping
                parse_inline_styles(p, text)
                for run in p.runs:
                    run.font.size = Pt(9.0)
                    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def process_code_block(doc, code_lines):
    if not code_lines:
        return
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    
    set_cell_background(cell, "F9FAFB")  # Subtle light gray
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    
    code_text = "".join(code_lines).rstrip()
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.0)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Default Normal style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    
    # Load markdown file
    with open("Web_Technology_Report.md", "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    state = "normal"
    buffer = []
    is_cover = True
    
    for line in lines:
        line_stripped = line.strip()
        
        # 1. State: Code block
        if state == "code_block":
            if line_stripped.startswith("```"):
                process_code_block(doc, buffer)
                buffer.clear()
                state = "normal"
            else:
                buffer.append(line)
            continue
            
        # 2. State: Table
        if state == "table":
            if line_stripped.startswith("|"):
                buffer.append(line)
            else:
                process_table(doc, buffer)
                buffer.clear()
                state = "normal"
                # Fall through to normal processing for the current line
            if state != "normal":
                continue
                
        # 3. State: Normal
        if line_stripped.startswith("```"):
            flush_paragraph_buffer(doc, buffer)
            state = "code_block"
            continue
            
        if line_stripped.startswith("|"):
            flush_paragraph_buffer(doc, buffer)
            state = "table"
            buffer.append(line)
            continue
            
        # Heading 1
        if line_stripped.startswith("# "):
            flush_paragraph_buffer(doc, buffer)
            heading_text = line_stripped[2:]
            add_heading(doc, heading_text, 1, is_cover=is_cover)
            continue
            
        # Heading 2
        if line_stripped.startswith("## "):
            flush_paragraph_buffer(doc, buffer)
            heading_text = line_stripped[3:]
            add_heading(doc, heading_text, 2, is_cover=is_cover)
            continue
            
        # Heading 3
        if line_stripped.startswith("### "):
            flush_paragraph_buffer(doc, buffer)
            heading_text = line_stripped[4:]
            add_heading(doc, heading_text, 3, is_cover=is_cover)
            continue
            
        # Heading 4
        if line_stripped.startswith("#### "):
            flush_paragraph_buffer(doc, buffer)
            heading_text = line_stripped[5:]
            add_heading(doc, heading_text, 4, is_cover=is_cover)
            continue
            
        # Page breaks
        if "page-break-after: always;" in line_stripped or "<div style=\"page-break-after: always;\">" in line_stripped:
            flush_paragraph_buffer(doc, buffer)
            doc.add_page_break()
            is_cover = False  # Cover page ends after first page break
            continue
            
        # Horizontal rule
        if line_stripped == "---":
            flush_paragraph_buffer(doc, buffer)
            add_horizontal_line(doc)
            continue
            
        # Lists
        if line_stripped.startswith("- ") or line_stripped.startswith("* "):
            flush_paragraph_buffer(doc, buffer)
            add_bullet_item(doc, line_stripped)
            continue
            
        if re.match(r'^\d+\.\s+', line_stripped):
            flush_paragraph_buffer(doc, buffer)
            add_numbered_item(doc, line_stripped)
            continue
            
        # Blank line
        if not line_stripped:
            flush_paragraph_buffer(doc, buffer)
            continue
            
        # Regular text
        buffer.append(line)
        
    # End of file cleanup
    if state == "code_block":
        process_code_block(doc, buffer)
    elif state == "table":
        process_table(doc, buffer)
    else:
        flush_paragraph_buffer(doc, buffer)
        
    # Save the document
    doc.save("Web_Technology_Report.docx")
    print("Web_Technology_Report.docx created successfully!")

if __name__ == "__main__":
    main()
